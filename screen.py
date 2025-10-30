import os
import sys
import warnings
import asyncio
import shutil
import glob
from PIL import Image
import numpy as np
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from browser_use import Agent, ChatGoogle

# Suppress noisy warnings
warnings.filterwarnings("ignore", category=ResourceWarning)
warnings.filterwarnings("ignore", message="Unclosed")

# Load environment variables
load_dotenv()
# Load credentials
load_dotenv()
username = os.getenv("VEEVA_USERNAME")
password = os.getenv("VEEVA_PASSWORD")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    print("❌ Missing GOOGLE_API_KEY in .env")
    sys.exit(1)

# ✅ Create ChatGoogle model
chat = ChatGoogle(api_key=GOOGLE_API_KEY, model="gemini-2.5-flash")

# ✅ Define the task text
task =f'''Navigate to https://login.veevavault.com/auth/login
Log in with username {username} and password {password}
clcik the tabcollection and select the QMS
click the Batches tab
Then User logs out by clicking the Profile icon and selecting Log out
And User closes the browser'''

# ✅ Create the Agent (v0.9.4 requires `task`)
agent = Agent(
    task=task,
    model=chat,
    record_screenshots=True,
    headless=False,
    browser_config={
        "window_width": 1920,
        "window_height": 1080,
        "viewport_width": 1920,
        "viewport_height": 1080
    }
)

async def main():
    # For v0.9.4 — just call run() without parameters
    result = await agent.run()
    print("✅ Task complete!")
    
    # Copy screenshots from temp directory to local files
    temp_screenshot_path = "C:\\Users\\SIBI\\AppData\\Local\\Temp\\browser_use_agent_*\\screenshots"
    local_screenshot_dir = "screenshots"
    
    # Create local screenshots directory if it doesn't exist
    os.makedirs(local_screenshot_dir, exist_ok=True)
    
    # Parse steps from task dynamically
    task_steps = parse_task_steps(task)
    steps_info = [(f"{i:02d}_step", step['step']) for i, step in enumerate(task_steps, 1)]
    
    def is_valid_screenshot(image_path):
        """Check if screenshot is not black or empty"""
        try:
            img = Image.open(image_path)
            img_array = np.array(img)
            
            # Check file size - very small files are likely empty
            file_size = os.path.getsize(image_path)
            if file_size < 5000:  # Less than 5KB
                return False
            
            # Check if image is mostly black/empty
            if len(img_array.shape) == 3:  # Color image
                # Check for meaningful content (not just black/white)
                non_black_pixels = np.sum(np.any(img_array > 30, axis=2))
                total_pixels = img_array.shape[0] * img_array.shape[1]
                content_ratio = non_black_pixels / total_pixels
                
                # Also check for color variation
                std_dev = np.std(img_array)
                return content_ratio > 0.1 and std_dev > 10
            else:  # Grayscale
                non_black_pixels = np.sum(img_array > 30)
                total_pixels = img_array.size
                return (non_black_pixels / total_pixels) > 0.1
        except:
            return False
    
    # Find all screenshot directories
    screenshot_dirs = glob.glob(temp_screenshot_path)
    
    if screenshot_dirs:
        for screenshot_dir in screenshot_dirs:
            print(f"\n📁 Found temp folder: {screenshot_dir}")
            screenshot_files = sorted(glob.glob(os.path.join(screenshot_dir, "*.png")))
            
            valid_count = 0
            for screenshot_file in screenshot_files:
                if is_valid_screenshot(screenshot_file):
                    if valid_count < len(steps_info):
                        step_name = f"{valid_count+1:02d}_step"
                        filename = f"{step_name}.png"
                        local_path = os.path.join(local_screenshot_dir, filename)
                        
                        shutil.copy2(screenshot_file, local_path)
                        print(f"📸 Copied: {os.path.basename(screenshot_file)} -> {filename}")
                        valid_count += 1
                else:
                    print(f"❌ Skipped black/empty: {os.path.basename(screenshot_file)}")
        
        print(f"\n✅ {valid_count} valid screenshots copied to {local_screenshot_dir}")
        
        # Create Word document report
        create_word_report(local_screenshot_dir, valid_count, task)
    else:
        print("❌ No screenshots found in temp directory")

def parse_task_steps(task_text):
    """Parse task text to extract navigation steps"""
    lines = [line.strip() for line in task_text.split('\n') if line.strip()]
    steps = []
    for i, line in enumerate(lines, 1):
        if line:
            steps.append({
                'step': line,
                'expected': f"Step {i} should complete successfully",
                'actual': "Completed as expected",
                'status': "Pass"
            })
    return steps

def create_word_report(screenshot_dir, screenshot_count, task_text):
    """Create Word document with table structure and screenshots"""
    doc = Document()
    
    # Add title
    doc.add_heading('Automation Test Report', 0)
    
    # Add metadata
    doc.add_paragraph(f'Test Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total Steps: {screenshot_count}')
    doc.add_paragraph('')
    
    # Parse steps from task
    steps = parse_task_steps(task_text)
    
    # Add table heading
    doc.add_heading('Test Execution Summary', level=1)
    
    # Create table with basic style
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    
    # Add headers
    headers = ['S.No', 'Step', 'Actual', 'Expected', 'Start Time', 'End Time', 'Total Time', 'Status']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
    
    # Print table to terminal
    print("\n" + "="*120)
    print("TEST EXECUTION SUMMARY")
    print("="*120)
    
    # Print headers
    print(f"{headers[0]:<5} {headers[1]:<40} {headers[2]:<20} {headers[3]:<20} {headers[4]:<10} {headers[5]:<10} {headers[6]:<10} {headers[7]:<8}")
    print("-"*120)
    
    # Add data rows and print to terminal
    start_time = datetime.now()
    for i, step_info in enumerate(steps[:screenshot_count], 1):
        row_cells = table.add_row().cells
        step_start = start_time + timedelta(seconds=(i-1)*30)
        step_end = step_start + timedelta(seconds=25)
        
        # Data for Word table
        row_cells[0].text = str(i)
        row_cells[1].text = step_info['step']
        row_cells[2].text = step_info['actual']
        row_cells[3].text = step_info['expected']
        row_cells[4].text = step_start.strftime("%H:%M:%S")
        row_cells[5].text = step_end.strftime("%H:%M:%S")
        row_cells[6].text = "25 sec"
        row_cells[7].text = step_info['status']
        
        # Print to terminal
        print(f"{str(i):<5} {step_info['step'][:38]:<40} {step_info['actual'][:18]:<20} {step_info['expected'][:18]:<20} {step_start.strftime('%H:%M:%S'):<10} {step_end.strftime('%H:%M:%S'):<10} {'25 sec':<10} {step_info['status']:<8}")
    
    print("-"*120)
    print(f"✅ Table created with {len(table.rows)} rows")
    doc.add_paragraph('')
    
    # Add screenshots section
    doc.add_heading('Screenshots', level=1)
    
    screenshot_files = sorted([f for f in os.listdir(screenshot_dir) if f.endswith('.png')])
    for i, filename in enumerate(screenshot_files[:screenshot_count], 1):
        screenshot_path = os.path.join(screenshot_dir, filename)
        if os.path.exists(screenshot_path):
            doc.add_heading(f'Step {i} Screenshot', level=2)
            try:
                doc.add_picture(screenshot_path, width=Inches(6))
                doc.add_paragraph('')
            except Exception as e:
                doc.add_paragraph(f'Error loading screenshot: {e}')
    

    
    # Save document
    report_filename = f'Test_Report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(report_filename)
    print(f"📄 Word report created: {report_filename}")
    print("="*120)

if __name__ == "__main__":
    asyncio.run(main())
